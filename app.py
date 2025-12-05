"""
API Flask para conversão de documentos DOC para PDF com substituição de tags

Autor: Maxwell da Silva Oliveira - M&S do Brasil LTDA
Email: maxwbh@gmail.com
LinkedIn: /maxwbh
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import base64
import io
import os
import tempfile
from docx import Document
import subprocess
from werkzeug.exceptions import BadRequest
import logging
import time
from datetime import datetime

# Importa informações de versão
from version import __version__, __author__, __email__, __company__, __linkedin__

app = Flask(__name__)
CORS(app)

# Configuração avançada de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - [%(funcName)s] - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Log de inicialização
logger.info(f"="*60)
logger.info(f"Iniciando DOC2PDF API v{__version__}")
logger.info(f"Desenvolvido por: {__author__} - {__company__}")
logger.info(f"="*60)


# Middleware para logging de requisições
@app.before_request
def log_request():
    """Log detalhado de cada requisição"""
    request.start_time = time.time()
    logger.info(f">>> NOVA REQUISIÇÃO")
    logger.info(f"Método: {request.method}")
    logger.info(f"Endpoint: {request.path}")
    logger.info(f"IP Cliente: {request.remote_addr}")
    logger.info(f"User-Agent: {request.headers.get('User-Agent', 'N/A')[:100]}")

    if request.method in ['POST', 'PUT', 'PATCH']:
        content_length = request.headers.get('Content-Length', 0)
        logger.info(f"Tamanho do payload: {content_length} bytes")


@app.after_request
def log_response(response):
    """Log da resposta e tempo de processamento"""
    if hasattr(request, 'start_time'):
        duration = time.time() - request.start_time
        logger.info(f"<<< RESPOSTA ENVIADA")
        logger.info(f"Status: {response.status_code}")
        logger.info(f"Tempo de processamento: {duration:.3f}s")
        logger.info(f"Tamanho da resposta: {response.content_length or 0} bytes")
        logger.info(f"-"*60)
    return response


# Headers de segurança
@app.after_request
def add_security_headers(response):
    """Adiciona headers de segurança em todas as respostas"""
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response


def decode_base64_file(base64_string):
    """Decodifica string Base64 para bytes"""
    try:
        return base64.b64decode(base64_string)
    except Exception as e:
        logger.error(f"Erro ao decodificar Base64: {str(e)}")
        raise ValueError("String Base64 inválida")


def validate_docx_format(doc_bytes):
    """
    Valida se os bytes são de um arquivo DOCX válido

    Args:
        doc_bytes: Bytes do documento

    Returns:
        tuple: (is_valid, error_message)
    """
    # DOCX é um arquivo ZIP que começa com PK (50 4B)
    if len(doc_bytes) < 4:
        return False, "Arquivo muito pequeno para ser um DOCX válido"

    # Verifica assinatura ZIP (DOCX files)
    if doc_bytes[0:2] == b'PK':
        logger.info("✓ Formato DOCX detectado (arquivo ZIP)")
        return True, None

    # Verifica assinatura de arquivo DOC antigo (D0 CF 11 E0)
    if doc_bytes[0:4] == b'\xD0\xCF\x11\xE0':
        logger.warning("⚠ Arquivo .DOC (formato antigo) detectado")
        return False, "Formato .DOC (Word 97-2003) não suportado. Por favor, converta para .DOCX primeiro"

    # Verifica se começa com texto (possível erro de encoding)
    if doc_bytes[0:5].decode('utf-8', errors='ignore').isprintable():
        logger.warning("⚠ Arquivo parece ser texto puro")
        return False, "Arquivo parece não ser um documento Word. Certifique-se de enviar um arquivo .DOCX válido em Base64"

    # Formato desconhecido
    logger.warning(f"⚠ Formato desconhecido. Primeiros bytes: {doc_bytes[0:4].hex()}")
    return False, "Formato de arquivo não reconhecido. Apenas arquivos .DOCX (Word 2007+) são suportados"


def replace_tags_in_doc(doc_bytes, replacements):
    """
    Substitui tags no documento Word pelos valores fornecidos

    Args:
        doc_bytes: Bytes do documento Word
        replacements: Dicionário com as substituições {tag: valor}

    Returns:
        Document: Documento modificado
    """
    try:
        # Valida o formato do documento antes de processar
        is_valid, error_msg = validate_docx_format(doc_bytes)
        if not is_valid:
            logger.error(f"Validação de formato falhou: {error_msg}")
            raise ValueError(error_msg)

        # Abre o documento a partir dos bytes
        doc = Document(io.BytesIO(doc_bytes))

        # Substitui tags nos parágrafos
        for paragraph in doc.paragraphs:
            for tag, value in replacements.items():
                # Formata a tag com %%
                tag_formatted = f"%%{tag.upper()}%%"
                if tag_formatted in paragraph.text:
                    # Substitui em cada run para preservar formatação
                    for run in paragraph.runs:
                        if tag_formatted in run.text:
                            run.text = run.text.replace(tag_formatted, str(value))

        # Substitui tags nas tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for tag, value in replacements.items():
                        tag_formatted = f"%%{tag.upper()}%%"
                        if tag_formatted in cell.text:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    if tag_formatted in run.text:
                                        run.text = run.text.replace(tag_formatted, str(value))

        return doc
    except Exception as e:
        logger.error(f"Erro ao processar documento: {str(e)}")
        raise ValueError(f"Erro ao processar documento Word: {str(e)}")


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Converte documento DOCX para PDF usando LibreOffice

    Args:
        docx_path: Caminho do arquivo DOCX
        pdf_path: Caminho onde salvar o PDF
    """
    try:
        # Usa LibreOffice para conversão (disponível no Docker)
        result = subprocess.run(
            [
                'libreoffice',
                '--headless',
                '--convert-to',
                'pdf',
                '--outdir',
                os.path.dirname(pdf_path),
                docx_path
            ],
            capture_output=True,
            text=True,
            timeout=30
        )

        if result.returncode != 0:
            logger.error(f"Erro LibreOffice: {result.stderr}")
            raise Exception(f"Erro na conversão para PDF: {result.stderr}")

        # LibreOffice salva com o mesmo nome base do arquivo de entrada
        generated_pdf = os.path.join(
            os.path.dirname(pdf_path),
            os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        )

        # Renomeia se necessário
        if generated_pdf != pdf_path and os.path.exists(generated_pdf):
            os.rename(generated_pdf, pdf_path)

    except subprocess.TimeoutExpired:
        raise Exception("Timeout na conversão do documento")
    except Exception as e:
        logger.error(f"Erro na conversão: {str(e)}")
        raise


@app.route('/health', methods=['GET'])
def health_check():
    """Endpoint para verificação de saúde da API"""
    return jsonify({
        'status': 'healthy',
        'service': 'doc2pdf-api',
        'version': __version__
    }), 200


@app.route('/convert', methods=['POST'])
def convert_document():
    """
    Endpoint principal para conversão de documentos

    Espera JSON com:
    - document: string Base64 do arquivo .DOC
    - replacements: objeto com as substituições {tag: valor}

    Retorna:
    - pdf: string Base64 do PDF gerado
    """
    try:
        # Valida requisição
        if not request.is_json:
            return jsonify({'error': 'Content-Type deve ser application/json'}), 400

        data = request.get_json()

        # Valida campos obrigatórios
        if 'document' not in data:
            return jsonify({'error': 'Campo "document" é obrigatório'}), 400

        if 'replacements' not in data:
            return jsonify({'error': 'Campo "replacements" é obrigatório'}), 400

        document_base64 = data['document']
        replacements = data['replacements']

        # Valida que replacements é um dicionário
        if not isinstance(replacements, dict):
            logger.warning("Requisição rejeitada: replacements inválido")
            return jsonify({'error': 'Campo "replacements" deve ser um objeto JSON'}), 400

        logger.info(f"✓ Validação OK - {len(replacements)} substituições encontradas")
        logger.info(f"Tags a substituir: {list(replacements.keys())}")

        # Decodifica o documento
        logger.info("Etapa 1/4: Decodificando documento Base64...")
        start_decode = time.time()
        doc_bytes = decode_base64_file(document_base64)
        logger.info(f"✓ Documento decodificado ({len(doc_bytes)} bytes) em {time.time() - start_decode:.3f}s")

        # Substitui as tags no documento
        logger.info("Etapa 2/4: Substituindo tags no documento...")
        start_replace = time.time()
        modified_doc = replace_tags_in_doc(doc_bytes, replacements)
        logger.info(f"✓ Tags substituídas em {time.time() - start_replace:.3f}s")

        # Cria arquivos temporários
        with tempfile.TemporaryDirectory() as temp_dir:
            logger.info(f"Diretório temporário: {temp_dir}")
            docx_path = os.path.join(temp_dir, 'document.docx')
            pdf_path = os.path.join(temp_dir, 'document.pdf')

            # Salva o documento modificado
            logger.info("Etapa 3/4: Salvando documento DOCX...")
            start_save = time.time()
            modified_doc.save(docx_path)
            doc_size = os.path.getsize(docx_path)
            logger.info(f"✓ DOCX salvo ({doc_size} bytes) em {time.time() - start_save:.3f}s")

            # Converte para PDF
            logger.info("Etapa 4/4: Convertendo DOCX para PDF...")
            start_convert = time.time()
            convert_docx_to_pdf(docx_path, pdf_path)

            # Verifica se o PDF foi gerado
            if not os.path.exists(pdf_path):
                logger.error("ERRO: PDF não foi gerado pelo LibreOffice")
                raise Exception("PDF não foi gerado")

            pdf_size = os.path.getsize(pdf_path)
            logger.info(f"✓ PDF gerado ({pdf_size} bytes) em {time.time() - start_convert:.3f}s")

            # Lê o PDF e converte para Base64
            logger.info("Codificando PDF para Base64...")
            start_encode = time.time()
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
                pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
            logger.info(f"✓ PDF codificado em Base64 ({len(pdf_base64)} chars) em {time.time() - start_encode:.3f}s")

        # Calcula tempo total
        total_time = time.time() - request.start_time if hasattr(request, 'start_time') else 0
        logger.info(f"✅ CONVERSÃO CONCLUÍDA COM SUCESSO")
        logger.info(f"Resumo: DOCX ({doc_size}b) -> PDF ({pdf_size}b) -> Base64 ({len(pdf_base64)} chars)")
        logger.info(f"Tempo total de conversão: {total_time:.3f}s")

        return jsonify({
            'success': True,
            'pdf': pdf_base64,
            'message': 'Documento convertido com sucesso'
        }), 200

    except ValueError as e:
        logger.error(f"Erro de validação: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Erro interno: {str(e)}")
        return jsonify({'error': f'Erro ao processar documento: {str(e)}'}), 500


@app.route('/convert-file', methods=['POST'])
def convert_document_file():
    """
    Endpoint para conversão de documentos retornando arquivo PDF

    Espera JSON com:
    - document: string Base64 do arquivo .DOC
    - replacements: objeto com as substituições {tag: valor}
    - filename (opcional): nome do arquivo PDF a ser gerado

    Retorna:
    - Arquivo PDF para download/visualização
    """
    temp_dir = None
    try:
        # Valida requisição
        if not request.is_json:
            return jsonify({'error': 'Content-Type deve ser application/json'}), 400

        data = request.get_json()

        # Valida campos obrigatórios
        if 'document' not in data:
            return jsonify({'error': 'Campo "document" é obrigatório'}), 400

        if 'replacements' not in data:
            return jsonify({'error': 'Campo "replacements" é obrigatório'}), 400

        document_base64 = data['document']
        replacements = data['replacements']
        filename = data.get('filename', 'documento.pdf')

        # Garante que o filename termina com .pdf
        if not filename.endswith('.pdf'):
            filename += '.pdf'

        # Valida que replacements é um dicionário
        if not isinstance(replacements, dict):
            logger.warning("Requisição rejeitada: replacements inválido")
            return jsonify({'error': 'Campo "replacements" deve ser um objeto JSON'}), 400

        logger.info(f"✓ Validação OK - {len(replacements)} substituições encontradas")
        logger.info(f"Tags a substituir: {list(replacements.keys())}")
        logger.info(f"Nome do arquivo de saída: {filename}")

        # Decodifica o documento
        logger.info("Etapa 1/4: Decodificando documento Base64...")
        start_decode = time.time()
        doc_bytes = decode_base64_file(document_base64)
        logger.info(f"✓ Documento decodificado ({len(doc_bytes)} bytes) em {time.time() - start_decode:.3f}s")

        # Substitui as tags no documento
        logger.info("Etapa 2/4: Substituindo tags no documento...")
        start_replace = time.time()
        modified_doc = replace_tags_in_doc(doc_bytes, replacements)
        logger.info(f"✓ Tags substituídas em {time.time() - start_replace:.3f}s")

        # Cria diretório temporário
        temp_dir = tempfile.mkdtemp()
        logger.info(f"Diretório temporário: {temp_dir}")
        docx_path = os.path.join(temp_dir, 'document.docx')
        pdf_path = os.path.join(temp_dir, 'document.pdf')

        # Salva o documento modificado
        logger.info("Etapa 3/4: Salvando documento DOCX...")
        start_save = time.time()
        modified_doc.save(docx_path)
        doc_size = os.path.getsize(docx_path)
        logger.info(f"✓ DOCX salvo ({doc_size} bytes) em {time.time() - start_save:.3f}s")

        # Converte para PDF
        logger.info("Etapa 4/4: Convertendo DOCX para PDF...")
        start_convert = time.time()
        convert_docx_to_pdf(docx_path, pdf_path)

        # Verifica se o PDF foi gerado
        if not os.path.exists(pdf_path):
            logger.error("ERRO: PDF não foi gerado pelo LibreOffice")
            raise Exception("PDF não foi gerado")

        pdf_size = os.path.getsize(pdf_path)
        logger.info(f"✓ PDF gerado ({pdf_size} bytes) em {time.time() - start_convert:.3f}s")

        # Calcula tempo total
        total_time = time.time() - request.start_time if hasattr(request, 'start_time') else 0
        logger.info(f"✅ CONVERSÃO CONCLUÍDA COM SUCESSO")
        logger.info(f"Resumo: DOCX ({doc_size}b) -> PDF ({pdf_size}b)")
        logger.info(f"Tempo total de conversão: {total_time:.3f}s")
        logger.info(f"Retornando arquivo: {filename}")

        # Retorna o arquivo PDF
        return send_file(
            pdf_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except ValueError as e:
        logger.error(f"Erro de validação: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Erro interno: {str(e)}")
        return jsonify({'error': f'Erro ao processar documento: {str(e)}'}), 500
    finally:
        # Limpa arquivos temporários após o envio
        if temp_dir and os.path.exists(temp_dir):
            try:
                import shutil
                shutil.rmtree(temp_dir)
            except Exception as e:
                logger.warning(f"Não foi possível remover diretório temporário: {str(e)}")


@app.route('/process', methods=['POST'])
def process_document():
    """
    Endpoint flexível para processamento de documentos

    Suporta múltiplos formatos de entrada e saída

    Args:
        input_type: 'base64' ou 'doc' (padrão: 'base64')
        document: Documento (Base64 ou bytes)
        output_type: 'pdf', 'doc', 'base64_pdf', 'base64_doc' (padrão: 'pdf')
        replacements: Dict com substituições de tags
        filename: Nome do arquivo de saída (opcional)

    Returns:
        Documento no formato especificado
    """
    temp_dir = None
    try:
        # Valida requisição
        if not request.is_json:
            return jsonify({'error': 'Content-Type deve ser application/json'}), 400

        data = request.get_json()

        # Valida campos obrigatórios
        if 'document' not in data:
            return jsonify({'error': 'Campo "document" é obrigatório'}), 400

        if 'replacements' not in data:
            return jsonify({'error': 'Campo "replacements" é obrigatório'}), 400

        # Parâmetros de processamento
        input_type = data.get('input_type', 'base64').lower()
        output_type = data.get('output_type', 'pdf').lower()
        document_data = data['document']
        replacements = data['replacements']
        filename = data.get('filename', 'documento')

        # Valida tipos
        valid_input_types = ['base64', 'doc']
        valid_output_types = ['pdf', 'doc', 'base64_pdf', 'base64_doc']

        if input_type not in valid_input_types:
            return jsonify({
                'error': f'input_type inválido. Use: {", ".join(valid_input_types)}'
            }), 400

        if output_type not in valid_output_types:
            return jsonify({
                'error': f'output_type inválido. Use: {", ".join(valid_output_types)}'
            }), 400

        if not isinstance(replacements, dict):
            logger.warning("Requisição rejeitada: replacements inválido")
            return jsonify({'error': 'Campo "replacements" deve ser um objeto JSON'}), 400

        logger.info(f"✓ Validação OK")
        logger.info(f"Configuração: input_type={input_type}, output_type={output_type}")
        logger.info(f"Substituições: {len(replacements)} tags - {list(replacements.keys())}")
        logger.info(f"Nome do arquivo: {filename}")

        # Processa documento de entrada
        logger.info(f"Etapa 1/4: Processando entrada ({input_type})...")
        start_input = time.time()
        if input_type == 'base64':
            doc_bytes = decode_base64_file(document_data)
            logger.info(f"✓ Base64 decodificado ({len(doc_bytes)} bytes) em {time.time() - start_input:.3f}s")
        else:  # doc
            doc_bytes = document_data.encode() if isinstance(document_data, str) else document_data
            logger.info(f"✓ Documento processado ({len(doc_bytes)} bytes) em {time.time() - start_input:.3f}s")

        # Substitui tags no documento
        logger.info("Etapa 2/4: Substituindo tags no documento...")
        start_replace = time.time()
        modified_doc = replace_tags_in_doc(doc_bytes, replacements)
        logger.info(f"✓ Tags substituídas em {time.time() - start_replace:.3f}s")

        # Cria diretório temporário
        temp_dir = tempfile.mkdtemp()
        logger.info(f"Diretório temporário: {temp_dir}")
        docx_path = os.path.join(temp_dir, 'document.docx')

        logger.info("Etapa 3/4: Salvando documento DOCX...")
        start_save = time.time()
        modified_doc.save(docx_path)
        doc_size = os.path.getsize(docx_path)
        logger.info(f"✓ DOCX salvo ({doc_size} bytes) em {time.time() - start_save:.3f}s")

        # Processa saída baseado no tipo solicitado
        if output_type == 'pdf':
            # Retorna arquivo PDF
            logger.info("Etapa 4/4: Convertendo para PDF e retornando arquivo...")
            start_output = time.time()
            pdf_path = os.path.join(temp_dir, 'document.pdf')
            convert_docx_to_pdf(docx_path, pdf_path)

            if not os.path.exists(pdf_path):
                logger.error("ERRO: PDF não foi gerado pelo LibreOffice")
                raise Exception("PDF não foi gerado")

            pdf_size = os.path.getsize(pdf_path)
            output_filename = filename if filename.endswith('.pdf') else f"{filename}.pdf"
            logger.info(f"✓ PDF gerado ({pdf_size} bytes) em {time.time() - start_output:.3f}s")

            total_time = time.time() - request.start_time if hasattr(request, 'start_time') else 0
            logger.info(f"✅ PROCESSAMENTO CONCLUÍDO COM SUCESSO")
            logger.info(f"Resumo: DOCX ({doc_size}b) -> PDF ({pdf_size}b)")
            logger.info(f"Tempo total: {total_time:.3f}s")
            logger.info(f"Retornando arquivo: {output_filename}")

            return send_file(
                pdf_path,
                mimetype='application/pdf',
                as_attachment=True,
                download_name=output_filename
            )

        elif output_type == 'doc':
            # Retorna arquivo DOC
            logger.info("Etapa 4/4: Preparando arquivo DOCX para retorno...")
            start_output = time.time()
            output_filename = filename if filename.endswith('.docx') else f"{filename}.docx"
            logger.info(f"✓ Arquivo DOCX pronto ({doc_size} bytes) em {time.time() - start_output:.3f}s")

            total_time = time.time() - request.start_time if hasattr(request, 'start_time') else 0
            logger.info(f"✅ PROCESSAMENTO CONCLUÍDO COM SUCESSO")
            logger.info(f"Resumo: DOCX ({doc_size}b)")
            logger.info(f"Tempo total: {total_time:.3f}s")
            logger.info(f"Retornando arquivo: {output_filename}")

            return send_file(
                docx_path,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=output_filename
            )

        elif output_type == 'base64_pdf':
            # Retorna PDF em Base64
            logger.info("Etapa 4/4: Convertendo para PDF e codificando em Base64...")
            start_output = time.time()
            pdf_path = os.path.join(temp_dir, 'document.pdf')

            logger.info("Sub-etapa 4a: Convertendo DOCX para PDF...")
            start_convert = time.time()
            convert_docx_to_pdf(docx_path, pdf_path)

            if not os.path.exists(pdf_path):
                logger.error("ERRO: PDF não foi gerado pelo LibreOffice")
                raise Exception("PDF não foi gerado")

            pdf_size = os.path.getsize(pdf_path)
            logger.info(f"✓ PDF gerado ({pdf_size} bytes) em {time.time() - start_convert:.3f}s")

            logger.info("Sub-etapa 4b: Codificando PDF para Base64...")
            start_encode = time.time()
            with open(pdf_path, 'rb') as pdf_file:
                pdf_bytes = pdf_file.read()
                pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
            logger.info(f"✓ PDF codificado em Base64 ({len(pdf_base64)} chars) em {time.time() - start_encode:.3f}s")

            total_time = time.time() - request.start_time if hasattr(request, 'start_time') else 0
            logger.info(f"✅ PROCESSAMENTO CONCLUÍDO COM SUCESSO")
            logger.info(f"Resumo: DOCX ({doc_size}b) -> PDF ({pdf_size}b) -> Base64 ({len(pdf_base64)} chars)")
            logger.info(f"Tempo total: {total_time:.3f}s")

            return jsonify({
                'success': True,
                'output_type': 'base64_pdf',
                'document': pdf_base64,
                'filename': f"{filename}.pdf",
                'size_bytes': len(pdf_bytes),
                'message': 'Documento processado com sucesso'
            }), 200

        elif output_type == 'base64_doc':
            # Retorna DOC em Base64
            logger.info("Etapa 4/4: Codificando DOCX para Base64...")
            start_output = time.time()
            with open(docx_path, 'rb') as doc_file:
                doc_bytes_output = doc_file.read()
                doc_base64 = base64.b64encode(doc_bytes_output).decode('utf-8')
            logger.info(f"✓ DOCX codificado em Base64 ({len(doc_base64)} chars) em {time.time() - start_output:.3f}s")

            total_time = time.time() - request.start_time if hasattr(request, 'start_time') else 0
            logger.info(f"✅ PROCESSAMENTO CONCLUÍDO COM SUCESSO")
            logger.info(f"Resumo: DOCX ({doc_size}b) -> Base64 ({len(doc_base64)} chars)")
            logger.info(f"Tempo total: {total_time:.3f}s")

            return jsonify({
                'success': True,
                'output_type': 'base64_doc',
                'document': doc_base64,
                'filename': f"{filename}.docx",
                'size_bytes': len(doc_bytes_output),
                'message': 'Documento processado com sucesso'
            }), 200

    except ValueError as e:
        logger.error(f"Erro de validação: {str(e)}")
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Erro interno: {str(e)}")
        return jsonify({'error': f'Erro ao processar documento: {str(e)}'}), 500
    finally:
        # Limpa arquivos temporários
        if temp_dir and os.path.exists(temp_dir):
            try:
                import shutil
                shutil.rmtree(temp_dir)
            except Exception as e:
                logger.warning(f"Não foi possível remover diretório temporário: {str(e)}")


@app.route('/', methods=['GET'])
def index():
    """Endpoint raiz com informações da API"""
    return jsonify({
        'service': 'DOC to PDF Converter API',
        'version': __version__,
        'author': f'{__author__} - {__company__}',
        'email': __email__,
        'linkedin': __linkedin__,
        'endpoints': {
            '/health': 'GET - Health check',
            '/convert': 'POST - Convert DOC to PDF with tag replacement (returns Base64)',
            '/convert-file': 'POST - Convert DOC to PDF with tag replacement (returns PDF file)',
            '/process': 'POST - Flexible document processing (supports multiple input/output formats)'
        },
        'usage': {
            'method': 'POST',
            'endpoint': '/convert',
            'content_type': 'application/json',
            'body': {
                'document': 'Base64 encoded .DOC file',
                'replacements': {
                    'NOME': 'Jose da Silva',
                    'ENDERECO': 'Rua qualquer coisa, Nro1, Bairro das colinas, São Paulo/SP - CEP: 48.4839-877',
                    'DATANASCIMENTO': '01/01/1990'
                }
            },
            'response': {
                'success': True,
                'pdf': 'Base64 encoded PDF file',
                'message': 'Documento convertido com sucesso'
            }
        }
    }), 200


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
