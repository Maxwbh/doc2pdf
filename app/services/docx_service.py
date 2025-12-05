"""
Serviço de manipulação de documentos DOCX

Autor: Maxwell da Silva Oliveira - M&S do Brasil LTDA
"""
import io
from typing import Dict
from docx import Document
from app.utils.logger import logger
from app.utils.validators import validate_docx_format
from config.settings import TAG_OPEN, TAG_CLOSE


class DocxService:
    """Serviço para manipulação de documentos DOCX"""

    @staticmethod
    def replace_tags_in_doc(doc_bytes: bytes, replacements: Dict[str, str]) -> Document:
        """
        Substitui tags no documento Word pelos valores fornecidos

        Tags no formato: {TAG}
        Substitui em: parágrafos, tabelas, cabeçalhos e rodapés

        Args:
            doc_bytes: Bytes do documento DOCX
            replacements: Dicionário com tags e valores {TAG: valor}

        Returns:
            Documento Word modificado (objeto Document)

        Raises:
            ValueError: Se o documento for inválido ou houver erro no processamento
        """
        try:
            logger.info(f"Iniciando substituição de tags no documento")
            logger.info(f"Tags a substituir: {list(replacements.keys())}")

            # Valida o formato do documento antes de processar
            is_valid, error_msg = validate_docx_format(doc_bytes)
            if not is_valid:
                logger.error(f"Validação de formato falhou: {error_msg}")
                raise ValueError(error_msg)

            # Abre o documento a partir dos bytes
            doc = Document(io.BytesIO(doc_bytes))

            tags_replaced_count = 0

            # Função auxiliar para substituir tags em runs
            def replace_in_runs(runs, tag, value):
                """Substitui tag em uma lista de runs, preservando formatação"""
                replaced = False
                tag_formatted = f"{TAG_OPEN}{tag.upper()}{TAG_CLOSE}"
                for run in runs:
                    if tag_formatted in run.text:
                        run.text = run.text.replace(tag_formatted, str(value))
                        replaced = True
                return replaced

            # 1. Substitui tags nos parágrafos principais
            logger.info("Substituindo tags nos parágrafos...")
            for paragraph in doc.paragraphs:
                for tag, value in replacements.items():
                    if replace_in_runs(paragraph.runs, tag, value):
                        tags_replaced_count += 1

            # 2. Substitui tags nas tabelas
            logger.info("Substituindo tags nas tabelas...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for tag, value in replacements.items():
                                if replace_in_runs(paragraph.runs, tag, value):
                                    tags_replaced_count += 1

            # 3. Substitui tags nos cabeçalhos (headers)
            logger.info("Substituindo tags nos cabeçalhos...")
            for section in doc.sections:
                # Header principal
                if section.header:
                    for paragraph in section.header.paragraphs:
                        for tag, value in replacements.items():
                            if replace_in_runs(paragraph.runs, tag, value):
                                tags_replaced_count += 1
                                logger.info(f"  ✓ Tag {TAG_OPEN}{tag.upper()}{TAG_CLOSE} substituída no cabeçalho")

                    # Tabelas no header
                    for table in section.header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for tag, value in replacements.items():
                                        if replace_in_runs(paragraph.runs, tag, value):
                                            tags_replaced_count += 1

            # 4. Substitui tags nos rodapés (footers)
            logger.info("Substituindo tags nos rodapés...")
            for section in doc.sections:
                # Footer principal
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        for tag, value in replacements.items():
                            if replace_in_runs(paragraph.runs, tag, value):
                                tags_replaced_count += 1
                                logger.info(f"  ✓ Tag {TAG_OPEN}{tag.upper()}{TAG_CLOSE} substituída no rodapé")

                    # Tabelas no footer
                    for table in section.footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for tag, value in replacements.items():
                                        if replace_in_runs(paragraph.runs, tag, value):
                                            tags_replaced_count += 1

            logger.info(f"✓ Substituição concluída: {tags_replaced_count} ocorrências substituídas")
            return doc

        except Exception as e:
            logger.error(f"Erro ao processar documento: {str(e)}")
            raise ValueError(f"Erro ao processar documento Word: {str(e)}")
